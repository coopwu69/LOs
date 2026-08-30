# -*- coding: utf-8 -*-
import os, sys, json, glob
import pdfplumber
import docx
import openpyxl

BASE = os.getcwd()

def extract_pdf(path):
    out = []
    try:
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                out.append(t)
    except Exception as e:
        out.append(f"[PDF ERROR: {e}]")
    return "\n".join(out)

def extract_docx(path):
    try:
        d = docx.Document(path)
        lines = []
        for para in d.paragraphs:
            lines.append(para.text)
        # also tables
        for ti, tbl in enumerate(d.tables):
            lines.append(f"\n[TABLE {ti+1}]")
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[DOCX ERROR: {e}]"

def extract_xlsx(path):
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"\n[SHEET: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[XLSX ERROR: {e}]"

results = {}
for root, dirs, files in os.walk(BASE):
    # skip .claude
    if '.claude' in root:
        continue
    for f in files:
        ext = os.path.splitext(f)[1].lower()
        if ext not in ('.pdf', '.docx', '.xlsx'):
            continue
        full = os.path.join(root, f)
        rel = os.path.relpath(full, BASE)
        if ext == '.pdf':
            txt = extract_pdf(full)
        elif ext == '.docx':
            txt = extract_docx(full)
        elif ext == '.xlsx':
            txt = extract_xlsx(full)
        results[rel] = txt
        print(f"EXTRACTED: {rel}  [{len(txt)} chars]")

with open(os.path.join(BASE, '_extracted.json'), 'w', encoding='utf-8') as fh:
    json.dump(results, fh, ensure_ascii=False, indent=2)
print(f"\nTOTAL FILES: {len(results)}")
print("Saved to _extracted.json")
