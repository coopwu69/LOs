"""Parse assessment source documents (.docx, .pdf, .xlsx) into structured data.

Supports 3 layouts:
  - rubric4: 2-col table, vertical 4-level options in cell (☑ ระดับดีมาก (...))
  - rubric5: 2-col table, vertical 5-level options in cell (5 (ยอดเยี่ยม): ...)
  - matrix:  6-col table, horizontal 4/3/2/1 columns, no descriptions

Auto-detects layout from table structure.
"""

import os
import re
import sys
import json
import unicodedata
from typing import Optional

from thai_repair import repair
from schema import DOMAIN_MAP, RUBRIC4_LABELS, RUBRIC5_LABELS, MATRIX4_LABELS, SCALE_STATUS

# ─── Text cleaning ───

PREFIX_PATTERNS = [
    re.compile(r"^\s*:\s*"),
    re.compile(r"^\s*(?:PLO|CLO|LO)\s*(\d+)\s*[:\.\)]\s*", re.IGNORECASE),
    re.compile(r"^\s*(\d+)\s*[\.\)]\s*"),  # "1. " or "1) "
]

CHECKBOX_RE = re.compile(r"[☑☐]")
# General rubric level names (4-level and 5-level)
LEVEL_NAMES = {
    "ดีเยี่ยม": 5, "ยอดเยี่ยม": 5,
    "ดีมาก": 4, "ดี": 3, "พอใช้": 2,
    "ต้องปรับปรุง": 1, "ควรปรับปรุง": 1, "ปรับปรุง": 1,
}
# rubric4: ☑ ระดับดีมาก (description) — also handles * ระดับยอดเยี่ยม (desc) and \uf06f ระดับดีเยี่ยม (desc)
RUBRIC_LEVEL_LINE_RE = re.compile(
    r"^\s*[\uf06f\uf0b7\u2022\u25cf\-\*☑☐]?\s*ระดับ(ดีเยี่ยม|ยอดเยี่ยม|ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ปรับปรุง)\s*\((.*)\)\s*$",
    re.MULTILINE
)
# Legacy: keep RUBRIC4_LINE_RE for backward compat
RUBRIC4_LINE_RE = re.compile(
    r"^\s*[☑☐]?\s*ระดับ(ดีมาก|ดี|พอใช้|ควรปรับปรุง)\s*\((.*)\)\s*$", re.MULTILINE
)
# rubric5: 5 (ยอดเยี่ยม): description (may have bullet prefix like \uf0b7)
RUBRIC5_LINE_RE = re.compile(
    r"^\s*[\uf0b7\u2022\u25cf\-\*]?\s*(\d)\s*\((.+?)\)\s*:\s*(.*)$", re.MULTILINE
)
# Horizontal rubric: ดีเยี่ยม (5 คะแนน) or ดีมาก (4 คะแนn) — label on first line, desc on next
HORIZONTAL_RUBRIC_LABEL_RE = re.compile(
    r"^\s*[\uf06f\uf0b7\u2022\u25cf\-\*]?\s*(ดีเยี่ยม|ยอดเยี่ยม|ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ปรับปรุง)\s*[\(（]\s*(\d+)\s*คะแนน\s*[\)）]\s*$",
    re.MULTILINE
)
# matrix section header: หัวข้อที่ 1: ความรู้ (Knowledge)\n(สอดคล้องกับ CLO1, CLO2)
MATRIX_SECTION_RE = re.compile(
    r"^\s*หัวข้อที่\s*(\d+)\s*:\s*(.+?)(?:\n|$)", re.MULTILINE
)
# matrix question: 1.1 text (first line only)
MATRIX_Q_RE = re.compile(
    r"^\s*(\d+\.\d+)\s+([^\n]+)"
)
# Simple numbered question: 1. text or 1. text (no sub-numbering)
SIMPLE_Q_RE = re.compile(
    r"^\s*(\d+)\s*[\.\)]\s*([^\n]+)"
)# domain header in rubric tables
DOMAIN_HEADER_RE = re.compile(
    r"^\s*(?:\(\d+\)\s*)?(?:ด้าน)?\s*(ความรู้|ทักษะ|จริยธรรม|ลักษณะบุคคล|ลักษณะส่วนบุคคล|ลักษณะ)\s*[\(\（]\s*(Knowledge|Skills|Ethics|Character)\s*[\)\）]\s*$",
    re.IGNORECASE
)
# LO code prefix
LO_CODE_RE = re.compile(r"^\s*((?:PLO|CLO|LO)\s*\d+)\s*[:\.\)]\s*", re.IGNORECASE)
# PLO line in preamble
PLO_LINE_RE = re.compile(r"^\s*((?:PLO|CLO|LO)\s*\d+)\s+(.+)$", re.IGNORECASE)
# Domain section header in preamble
PREAMBLE_DOMAIN_RE = re.compile(
    r"^\s*(?:\d+\)\s*)?(?:ด้าน)?\s*(ความรู้|ทักษะ|จริยธรรม|ลักษณะบุคคล|ลักษณะส่วนบุคคล|ลักษณะ)\s*[\(\（]\s*(Knowledge|Skills|Ethics|Character)\s*[\)\）]",
    re.IGNORECASE
)
# Course code like PAD67-491
COURSE_CODE_RE = re.compile(r"\b([A-Z]{2,5}\d{2}-\d{3})\b")
# Revision label
REVISION_RE = re.compile(r"(?:หลักสูตรปรับปรุง|ปรับปรุง)\s*(?:ปีการศึกษา|พ\.ศ\.?)\s*([๐-๙\d]+)")


def clean_text(text: str) -> tuple[str, list[str]]:
    """Clean a text string: strip prefixes, checkboxes, normalize, repair Thai."""
    if not text:
        return "", []
    text = unicodedata.normalize("NFC", text)
    # Remove checkboxes
    text = CHECKBOX_RE.sub("", text)
    # Strip leading ": "
    text = re.sub(r"^\s*:\s*", "", text)
    # Collapse whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    # Thai repair
    text, reps = repair(text)
    return text, reps


def extract_lo_code(text: str) -> tuple[str, str]:
    """Extract LO code prefix from text. Returns (lo_code, remaining_text)."""
    m = LO_CODE_RE.match(text)
    if m:
        code = m.group(1).replace(" ", "")
        rest = text[m.end():].strip()
        return code, rest
    return "", text


def extract_english(text: str) -> tuple[str, str]:
    """Extract English text from parentheses at end. Returns (thai, english)."""
    # Match (English text) at end
    m = re.search(r"\(([^()]+)\)\s*$", text)
    if m and re.search(r"[a-zA-Z]{3,}", m.group(1)):
        return text[:m.start()].strip(), m.group(1).strip()
    return text, ""


def map_domain(thai_domain: str, eng_domain: str = "") -> str:
    """Map domain text to canonical domain type."""
    key = thai_domain.strip()
    if key in DOMAIN_MAP:
        return DOMAIN_MAP[key]
    if eng_domain:
        ek = eng_domain.lower().strip()
        if ek in DOMAIN_MAP:
            return DOMAIN_MAP[ek]
    return "general"


def detect_part(domain: str) -> int:
    """Determine part (1=knowledge+skills, 2=ethics+character)."""
    if domain in ("knowledge", "skills"):
        return 1
    if domain in ("ethics", "character"):
        return 2
    return 1


# ─── Layout detection ───

def detect_layout(tables: list) -> str:
    """Auto-detect layout from table structures.
    tables: list of (rows, cols, rows_data) tuples.
    """
    # Check for matrix (5+ columns) first
    for n_rows, n_cols, _ in tables:
        if n_cols >= 5:
            return "matrix"
    # Check cell content for rubric4 vs rubric5 across ALL 2-col tables
    has_rubric5 = False
    has_rubric4 = False
    level_prefix_re = re.compile(
        r"[\uf06f\uf0b7\u2022\u25cf\-\*☑☐]?\s*ระดับ(ดีเยี่ยม|ยอดเยี่ยม|ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ปรับปรุง)"
    )
    for n_rows, n_cols, data in tables:
        if n_cols == 2 and n_rows > 2:
            for row in data:
                if len(row) >= 2 and row[1]:
                    cell = row[1]
                    if RUBRIC5_LINE_RE.search(cell):
                        has_rubric5 = True
                    if RUBRIC4_LINE_RE.search(cell):
                        has_rubric4 = True
                    # Check for level prefix (handles multi-line cells)
                    for m in level_prefix_re.finditer(cell):
                        level = m.group(1)
                        if LEVEL_NAMES.get(level, 0) >= 5:
                            has_rubric5 = True
                        else:
                            has_rubric4 = True
                    if HORIZONTAL_RUBRIC_LABEL_RE.search(cell):
                        for m in HORIZONTAL_RUBRIC_LABEL_RE.finditer(cell):
                            score = int(m.group(2))
                            if score >= 5:
                                has_rubric5 = True
                            else:
                                has_rubric4 = True
    if has_rubric5:
        return "rubric5"
    if has_rubric4:
        return "rubric4"
    # Default
    return "rubric4"


# ─── DOCX extraction ───

def extract_docx_tables(path: str) -> list:
    """Extract tables from .docx as list of (n_rows, n_cols, data)."""
    from docx import Document
    doc = Document(path)
    result = []
    for t in doc.tables:
        data = []
        for row in t.rows:
            data.append([cell.text for cell in row.cells])
        result.append((len(t.rows), len(t.columns), data))
    return result


def extract_docx_paragraphs(path: str) -> list[str]:
    """Extract paragraphs from .docx."""
    from docx import Document
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]


# ─── PDF extraction ───

def extract_pdf_tables(path: str) -> tuple[list, list[str]]:
    """Extract tables and text from PDF. Returns (tables, all_text_lines)."""
    import pdfplumber
    tables = []
    text_lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for t in page_tables:
                if t and len(t) > 1:
                    data = [[(c or "") for c in row] for row in t]
                    n_cols = max(len(r) for r in data) if data else 0
                    tables.append((len(data), n_cols, data))
            txt = page.extract_text() or ""
            text_lines.extend(txt.split("\n"))
    return tables, text_lines


# ─── Preamble parsing (PLOs, course codes, revision) ───

def parse_preamble(lines: list[str]) -> dict:
    """Parse preamble text for PLOs, course codes, revision label."""
    plos = []
    course_codes = set()
    revision_label = None
    current_domain = "general"
    plo_seq = 0

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Course codes
        for m in COURSE_CODE_RE.finditer(line):
            course_codes.add(m.group(1))

        # Revision label
        m = REVISION_RE.search(line)
        if m and not revision_label:
            year = m.group(1)
            # Convert Thai digits
            year = year.translate(str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789"))
            revision_label = f"หลักสูตรปรับปรุง ปีการศึกษา {year}"

        # Domain header
        m = PREAMBLE_DOMAIN_RE.match(line)
        if m:
            current_domain = map_domain(m.group(1), m.group(2))
            continue

        # PLO line
        m = PLO_LINE_RE.match(line)
        if m:
            code = m.group(1).replace(" ", "")
            text, reps = clean_text(m.group(2))
            plo_seq += 1
            plos.append({
                "code": code,
                "domain": current_domain,
                "text": text,
                "sequence": plo_seq,
            })

    return {
        "plos": plos,
        "course_codes": sorted(course_codes),
        "revision_label": revision_label,
    }


# ─── Rubric4 parser ───

def parse_rubric4_cell(cell_text: str) -> list[dict]:
    """Parse a rubric4 options cell (vertical 4-level with descriptions)."""
    options = []
    for line in cell_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = RUBRIC4_LINE_RE.match(line)
        if m:
            level, desc = m.group(1), m.group(2)
            score = {"ดีมาก": 4, "ดี": 3, "พอใช้": 2, "ควรปรับปรุง": 1}[level]
            label = f"ระดับ{level}"
            desc, _ = clean_text(desc)
            options.append({
                "score": score,
                "label_th": label,
                "description_th": desc,
            })
    # Sort by score desc and add sequence
    options.sort(key=lambda o: -o["score"])
    for i, o in enumerate(options):
        o["sequence"] = i + 1
    return options


def parse_rubric4_table(table_data: list) -> list[dict]:
    """Parse a rubric4 2-column table into sections + questions."""
    sections = []
    current_section = None
    q_seq = 0

    for row in table_data:
        if len(row) < 2:
            continue
        left = (row[0] or "").strip()
        right = (row[1] or "").strip()

        # Skip header row
        if "ผลลัพธ์การเรียนรู้" in left and "ผลการประเมิน" in right:
            continue

        # Domain header row (left has domain, right empty)
        m = DOMAIN_HEADER_RE.match(left)
        if m and not right:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(left)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # Question row (left has LO text, right has options)
        if left and right:
            lo_code, q_text = extract_lo_code(left)
            q_text, _ = clean_text(q_text)
            options = parse_rubric4_cell(right)
            if not options:
                continue
            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections


# ─── General rubric cell parser (handles all vertical variants) ───

def parse_general_rubric_cell(cell_text: str) -> list[dict]:
    """Parse a rubric options cell handling all vertical formats:
    - ☑ ระดับดีมาก (desc) — rubric4
    - * ระดับยอดเยี่ยม (desc) — CHI variant
    - \uf06f ระดับดีเยี่ยม (desc) — ENG variant (desc may span multiple lines)
    - 5 (ยอดเยี่ยม): desc — rubric5
    - ดีเยี่ยม (5 คะแนน)\ndesc — horizontal variant
    """
    options = []
    lines = cell_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Try rubric5 format: N (label): desc
        m5 = RUBRIC5_LINE_RE.match(line)
        if m5:
            score = int(m5.group(1))
            label = m5.group(2).strip()
            desc = m5.group(3).strip()
            j = i + 1
            while j < len(lines):
                nl = lines[j].strip()
                if not nl or RUBRIC5_LINE_RE.match(nl) or RUBRIC_LEVEL_LINE_RE.match(nl) or HORIZONTAL_RUBRIC_LABEL_RE.match(nl) or _is_level_start(nl):
                    break
                desc += " " + nl
                j += 1
            desc, _ = clean_text(desc)
            options.append({"score": score, "label_th": label, "description_th": desc})
            i = j
            continue

        # Try general level format: ระดับX (desc — may span multiple lines)
        # Match the prefix: ระดับX (
        ml_prefix = re.match(
            r"^\s*[\uf06f\uf0b7\u2022\u25cf\-\*☑☐]?\s*ระดับ(ดีเยี่ยม|ยอดเยี่ยม|ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ปรับปรุง)\s*[\(（]\s*(.*)$",
            line
        )
        if ml_prefix:
            level = ml_prefix.group(1)
            desc = ml_prefix.group(2)
            score = LEVEL_NAMES.get(level, 0)
            label = f"ระดับ{level}"
            # Collect continuation lines until closing ) or next level start
            j = i + 1
            closed = desc.endswith(")") or desc.endswith("）")
            if not closed:
                while j < len(lines):
                    nl = lines[j].strip()
                    if not nl:
                        j += 1
                        continue
                    if _is_level_start(nl) or RUBRIC5_LINE_RE.match(nl) or HORIZONTAL_RUBRIC_LABEL_RE.match(nl):
                        break
                    desc += " " + nl
                    if nl.endswith(")") or nl.endswith("）"):
                        j += 1
                        break
                    j += 1
            else:
                j = i + 1
            # Strip trailing )
            desc = re.sub(r"[\)）]\s*$", "", desc).strip()
            desc, _ = clean_text(desc)
            options.append({"score": score, "label_th": label, "description_th": desc})
            i = j
            continue

        # Try horizontal format: ดีเยี่ยม (5 คะแนน) — desc on next lines
        mh = HORIZONTAL_RUBRIC_LABEL_RE.match(line)
        if mh:
            level = mh.group(1)
            score = int(mh.group(2))
            label = f"ระดับ{level}" if not level.startswith("ระดับ") else level
            desc_parts = []
            j = i + 1
            while j < len(lines):
                nl = lines[j].strip()
                if not nl or RUBRIC5_LINE_RE.match(nl) or RUBRIC_LEVEL_LINE_RE.match(nl) or HORIZONTAL_RUBRIC_LABEL_RE.match(nl) or _is_level_start(nl):
                    break
                desc_parts.append(nl)
                j += 1
            desc = " ".join(desc_parts)
            desc, _ = clean_text(desc)
            options.append({"score": score, "label_th": label, "description_th": desc})
            i = j
            continue

        i += 1

    options.sort(key=lambda o: -o["score"])
    for idx, o in enumerate(options):
        o["sequence"] = idx + 1
    return options


def _is_level_start(line: str) -> bool:
    """Check if a line starts a new rubric level entry."""
    line = line.strip()
    if not line:
        return False
    if RUBRIC5_LINE_RE.match(line):
        return True
    if HORIZONTAL_RUBRIC_LABEL_RE.match(line):
        return True
    # Check for ระดับX prefix
    if re.match(r"^\s*[\uf06f\uf0b7\u2022\u25cf\-\*☑☐]?\s*ระดับ(ดีเยี่ยม|ยอดเยี่ยม|ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ปรับปรุง)", line):
        return True
    return False


def parse_general_rubric_table(table_data: list) -> list[dict]:
    """Parse a 2-column rubric table using the general cell parser."""
    sections = []
    current_section = None
    q_seq = 0

    for row in table_data:
        if len(row) < 2:
            continue
        left = (row[0] or "").strip()
        right = (row[1] or "").strip()

        # Skip header rows
        if ("ผลลัพธ์การเรียนรู้" in left or "ผลการเรียนรู้" in left) and "ผลการประเมิน" in right:
            continue
        if "หัวข้อประเมิน" in left and "คะแนนประเมิน" in right:
            continue
        if "ลักษณะบุคคล/สมรรถนะ" in left:
            continue

        # Domain header row
        m = DOMAIN_HEADER_RE.match(left)
        if m and not right:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(left)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # Question row
        if left and right:
            lo_code, q_text = extract_lo_code(left)
            if not lo_code:
                # Try simple numbered: "1. text"
                sm = SIMPLE_Q_RE.match(left)
                if sm:
                    lo_code = sm.group(1)
                    q_text = sm.group(2)
            q_text, _ = clean_text(q_text)
            options = parse_general_rubric_cell(right)
            if not options:
                continue
            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections


# ─── Rubric5 parser ───

def parse_rubric5_cell(cell_text: str) -> list[dict]:
    """Parse a rubric5 options cell (vertical 5-level with descriptions)."""
    options = []
    lines = cell_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        m = RUBRIC5_LINE_RE.match(line)
        if m:
            score = int(m.group(1))
            label = m.group(2).strip()
            desc = m.group(3).strip()
            # Description may continue on next lines until next "N (label):"
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line or RUBRIC5_LINE_RE.match(next_line):
                    break
                desc += " " + next_line
                j += 1
            desc, _ = clean_text(desc)
            options.append({
                "score": score,
                "label_th": label,
                "description_th": desc,
            })
            i = j
        else:
            i += 1
    options.sort(key=lambda o: -o["score"])
    for i, o in enumerate(options):
        o["sequence"] = i + 1
    return options


def parse_rubric5_table(table_data: list) -> list[dict]:
    """Parse a rubric5 2-column table."""
    sections = []
    current_section = None
    q_seq = 0

    for row in table_data:
        if len(row) < 2:
            continue
        left = (row[0] or "").strip()
        right = (row[1] or "").strip()

        if "ผลการเรียนรู้" in left and "ผลการประเมิน" in right:
            continue
        if "ผลลัพธ์การเรียนรู้" in left and "ผลการประเมิน" in right:
            continue

        m = DOMAIN_HEADER_RE.match(left)
        if m and not right:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(left)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        if left and right:
            lo_code, q_text = extract_lo_code(left)
            q_text, _ = clean_text(q_text)
            options = parse_rubric5_cell(right)
            if not options:
                continue
            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections


# ─── Matrix parser ───

def parse_matrix_table(table_data: list) -> list[dict]:
    """Parse a matrix 6-column table (Indicators | 4 | 3 | 2 | 1 | comments)."""
    sections = []
    current_section = None
    q_seq = 0

    for row in table_data:
        if not row or not row[0]:
            continue
        left = row[0].strip()

        # Skip header
        if "รายการประเมิน" in left or "Indicators" in left:
            continue

        # Section header: หัวข้อที่ N: domain (สอดคล้องกับ CLOx)
        full_text = left
        if len(row) > 1:
            # Sometimes CLO refs are in the next cell or same cell
            pass

        m = MATRIX_SECTION_RE.match(left)
        if m:
            # Extract domain from the title
            rest = m.group(2)
            # Look for (Knowledge) etc
            dm = re.search(r"(ความรู้|ทักษะ|จริยธรรม|ลักษณะบุคคล|ลักษณะส่วนบุคคล|ลักษณะ)\s*[\(\（]\s*(Knowledge|Skills|Ethics|Character)", rest, re.IGNORECASE)
            if dm:
                domain = map_domain(dm.group(1), dm.group(2))
                title = clean_text(rest)[0]
            else:
                domain = "general"
                title = clean_text(rest)[0]

            # Extract CLO refs from the full cell text (may be on next line)
            plo_refs = re.findall(r"(?:CLO|PLO|LO)\s*\d+", left, re.IGNORECASE)
            plo_refs = [r.replace(" ", "") for r in plo_refs]

            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
                "_plo_refs": plo_refs,
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # Question row: 1.1 text (English on next line)
        m = MATRIX_Q_RE.match(left)
        if m:
            lo_code = m.group(1)
            q_text = m.group(2)
            # English might be in parentheses on the same line or on the next line
            lines = left.split("\n")
            text_en = ""
            if len(lines) > 1:
                # English is usually in parentheses on the second line
                en_line = lines[1].strip()
                if en_line.startswith("(") and en_line.endswith(")"):
                    text_en = en_line[1:-1].strip()
                elif re.search(r"[a-zA-Z]{3,}", en_line):
                    text_en = en_line
            # Also check for English in parentheses on the first line
            q_text, en_from_paren = extract_english(q_text)
            if en_from_paren and not text_en:
                text_en = en_from_paren
            q_text, _ = clean_text(q_text)
            text_en, _ = clean_text(text_en)

            # Create 4 standard options (no descriptions for matrix)
            options = []
            for score, label in [(4, MATRIX4_LABELS[4]), (3, MATRIX4_LABELS[3]), (2, MATRIX4_LABELS[2]), (1, MATRIX4_LABELS[1])]:
                options.append({
                    "score": score,
                    "label_th": label,
                    "description_th": "",
                    "sequence": 5 - score,
                })

            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)

            plo_refs = current_section.get("_plo_refs", [])
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": text_en,
                "plo_refs": plo_refs,
                "sequence": q_seq,
                "options": options,
            })

    # Clean up _plo_refs
    for s in sections:
        s.pop("_plo_refs", None)

    return sections


# ─── Horizontal rubric4 (PDF variant: 5 columns) ───

def parse_horizontal_rubric4(table_data: list) -> list[dict]:
    """Parse a horizontal rubric4 table (question | desc4 | desc3 | desc2 | desc1).
    Used by some PDFs like ENVH.
    """
    sections = []
    current_section = None
    q_seq = 0

    for row in table_data:
        if not row or not row[0]:
            continue
        left = row[0].strip()

        # Skip header
        if "ผลการเรียนรู้" in left or "ผลลัพธ์การเรียนรู้" in left:
            continue

        # Domain header
        m = DOMAIN_HEADER_RE.match(left)
        if m:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(left)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # Question row: left=question, cols 1-4=descriptions
        if len(row) >= 5:
            lo_code, q_text = extract_lo_code(left)
            q_text, _ = clean_text(q_text)
            if not q_text or len(q_text) < 10:
                continue

            options = []
            for idx, score in [(1, 4), (2, 3), (3, 2), (4, 1)]:
                desc = (row[idx] or "").strip() if idx < len(row) else ""
                desc, _ = clean_text(desc)
                # Remove leading label like "ดีมาก (4 คะแนน)" from description
                desc = re.sub(r"^(?:ระดับ)?(?:ดีมาก|ดี|พอใช้|ต้องปรับปรุง|ควรปรับปรุง|ดีเยี่ยม)\s*[\(（].*?[\)）]\s*", "", desc).strip()
                options.append({
                    "score": score,
                    "label_th": RUBRIC4_LABELS[score],
                    "description_th": desc,
                    "sequence": 5 - score,
                })

            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections


# ─── Simple matrix parser (POL-style: numbered questions, no section headers) ───

def parse_simple_matrix_table(table_data: list, n_cols: int = 6) -> list[dict]:
    """Parse a simple matrix table with numbered questions and no section headers.
    POL format: คำถาม | 1 | 2 | 3 | 4 | 5 (5-level, no descriptions)
    """
    sections = []
    current_section = {
        "domain": "general",
        "title_th": "ข้อประเมินทั่วไป",
        "part": 1,
        "questions": [],
    }
    sections.append(current_section)
    q_seq = 0

    # Determine score columns from header
    header = table_data[0] if table_data else []
    score_cols = []
    for ci, cell in enumerate(header[1:], 1):
        val = (cell or "").strip()
        if val.isdigit():
            score_cols.append((ci, int(val)))

    for row in table_data[1:]:
        if not row or not row[0]:
            continue
        left = row[0].strip()
        if not left:
            continue

        # Try simple numbered: "1. text"
        m = SIMPLE_Q_RE.match(left)
        if m:
            lo_code = m.group(1)
            q_text = m.group(2)
            q_text, _ = clean_text(q_text)
            if len(q_text) < 10:
                continue

            # Determine number of levels from header
            n_levels = len(score_cols) if score_cols else 5
            options = []
            for idx, (ci, score) in enumerate(sorted(score_cols, key=lambda x: -x[1]) if score_cols else [(i, n_levels-i) for i in range(1, n_levels+1)]):
                label = RUBRIC5_LABELS.get(score, f"ระดับ{score}")
                options.append({
                    "score": score,
                    "label_th": label,
                    "description_th": "",
                    "sequence": idx + 1,
                })

            q_seq += 1
            current_section["questions"].append({
                "lo_code": lo_code,
                "text": q_text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections if current_section["questions"] else []


# ─── Paragraph-based parser (for FSI, CUL — no tables) ───

def parse_paragraph_sections(text_lines: list[str]) -> list[dict]:
    """Parse assessment data from plain text paragraphs when no tables are available.
    Looks for domain headers and LO/PLO/CLO lines.
    """
    sections = []
    current_section = None
    q_seq = 0

    for line in text_lines:
        line = line.strip()
        if not line:
            continue

        # Domain header
        m = DOMAIN_HEADER_RE.match(line)
        if m:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(line)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # PREAMBLE_DOMAIN_RE (without requiring end of line)
        m = PREAMBLE_DOMAIN_RE.match(line)
        if m:
            domain = map_domain(m.group(1), m.group(2))
            title = clean_text(line)[0]
            current_section = {
                "domain": domain,
                "title_th": title,
                "part": detect_part(domain),
                "questions": [],
            }
            sections.append(current_section)
            q_seq = 0
            continue

        # LO/CLO/PLO line as a question
        m = PLO_LINE_RE.match(line)
        if m:
            code = m.group(1).replace(" ", "")
            text, _ = clean_text(m.group(2))
            if len(text) < 10:
                continue
            q_seq += 1
            if current_section is None:
                current_section = {
                    "domain": "general",
                    "title_th": "ทั่วไป",
                    "part": 1,
                    "questions": [],
                }
                sections.append(current_section)
            # Create standard 4-level options (no descriptions)
            options = []
            for score, label in [(4, RUBRIC4_LABELS[4]), (3, RUBRIC4_LABELS[3]), (2, RUBRIC4_LABELS[2]), (1, RUBRIC4_LABELS[1])]:
                options.append({
                    "score": score,
                    "label_th": label,
                    "description_th": "",
                    "sequence": 5 - score,
                })
            current_section["questions"].append({
                "lo_code": code,
                "text": text,
                "text_en": "",
                "plo_refs": [],
                "sequence": q_seq,
                "options": options,
            })

    return sections


# ─── XLSX parser ───

def extract_xlsx_data(path: str) -> tuple[list, list[str]]:
    """Extract table-like data from XLSX. Returns (tables, text_lines)."""
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    tables = []
    text_lines = []

    for sn in wb.sheetnames:
        ws = wb[sn]
        # Convert sheet to table data
        data = []
        for r in range(1, ws.max_row + 1):
            row = []
            for c in range(1, min(ws.max_column + 1, 20)):
                val = ws.cell(r, c).value
                row.append(str(val) if val is not None else "")
            # Skip completely empty rows
            if any(cell.strip() for cell in row):
                data.append(row)
                # Also add to text lines
                for cell in row:
                    if cell.strip():
                        text_lines.append(cell.strip())
        if data:
            n_cols = max(len(r) for r in data)
            tables.append((len(data), n_cols, data))

    return tables, text_lines


# ─── Main parse function ───

def parse_document(path: str, layout_hint: str = None) -> dict:
    """Parse a source document. Returns structured data dict.

    Returns:
        {
            "plos": [...],
            "sections": [...],
            "course_codes": [...],
            "revision_label": str,
            "layout": str,  # detected
            "scale_status": str,
            "errors": [...],
            "thai_repairs": [...],
        }
    """
    ext = os.path.splitext(path)[1].lower()
    errors = []
    all_repairs = []

    if ext == ".docx":
        tables = extract_docx_tables(path)
        paragraphs = extract_docx_paragraphs(path)
        text_lines = paragraphs
    elif ext == ".pdf":
        tables, text_lines = extract_pdf_tables(path)
        paragraphs = text_lines
    elif ext == ".xlsx":
        tables, text_lines = extract_xlsx_data(path)
        paragraphs = text_lines
    else:
        return {
            "plos": [], "sections": [], "course_codes": [],
            "revision_label": None, "layout": "unknown",
            "scale_status": "standard_4", "errors": [f"unsupported ext: {ext}"],
            "thai_repairs": [],
        }

    # Parse preamble
    preamble = parse_preamble(text_lines)

    # Detect layout
    table_shapes = [(n, c, d) for n, c, d in tables]
    layout = detect_layout(table_shapes)

    # Override with hint if detection disagrees and hint is specific
    # But trust detection for rubric4 vs rubric5 since hints were often wrong
    if layout_hint == "matrix" and any(c >= 5 for _, c, _ in table_shapes):
        layout = "matrix"

    # Parse tables based on layout
    all_sections = []
    for n_rows, n_cols, data in tables:
        if n_cols >= 5:
            # Could be matrix, horizontal rubric4, or simple matrix
            header = data[0] if data else []
            header_text = " ".join((c or "") for c in header)

            # Check if it's a simple matrix (POL-style: "คำถาม | 1 | 2 | 3 | 4 | 5")
            if "คำถาม" in header_text or "คามถาม" in header_text:
                sections = parse_simple_matrix_table(data, n_cols)
                if sections:
                    layout = "matrix"
            # Check if it's horizontal rubric4 (question | desc4 | desc3 | desc2 | desc1)
            elif "4" in header_text and "3" in header_text and "2" in header_text and "1" in header_text:
                has_desc = any(
                    (row[1] or "").strip() if len(row) > 1 else ""
                    for row in data[1:5]
                    if row and len(row) > 1 and (row[1] or "").strip()
                )
                if has_desc and "Indicators" not in header_text and "รายการประเมิน" not in header_text:
                    sections = parse_horizontal_rubric4(data)
                    layout = "rubric4"
                else:
                    sections = parse_matrix_table(data)
                    layout = "matrix"
            else:
                # Try matrix first, then simple matrix
                sections = parse_matrix_table(data)
                if not sections:
                    sections = parse_simple_matrix_table(data, n_cols)
                layout = "matrix"
        elif n_cols == 2:
            # Use general rubric parser which handles all variants
            sections = parse_general_rubric_table(data)
            if not sections:
                # Fallback to specific parsers
                s4 = parse_rubric4_table(data)
                s5 = parse_rubric5_table(data)
                opts4 = sum(len(q["options"]) for s in s4 for q in s["questions"])
                opts5 = sum(len(q["options"]) for s in s5 for q in s["questions"])
                if opts5 > opts4:
                    sections = s5
                    if opts5 > 0:
                        layout = "rubric5"
                elif opts4 > 0:
                    sections = s4
                    layout = "rubric4"
            elif layout == "rubric4":
                # Check if general parser found 5-level options
                max_score = max((o["score"] for s in sections for q in s["questions"] for o in q["options"]), default=4)
                if max_score >= 5:
                    layout = "rubric5"
        else:
            continue

        if sections:
            all_sections.extend(sections)

    # Fallback: if no sections found from tables, try paragraph-based parsing
    if not all_sections and text_lines:
        para_sections = parse_paragraph_sections(text_lines)
        if para_sections:
            all_sections.extend(para_sections)
            if layout == "rubric4":
                layout = "rubric4"  # keep default
            errors.append("Used paragraph fallback (no table sections found)")

    # Collect Thai repairs from all text
    for line in text_lines:
        _, reps = clean_text(line)
        all_repairs.extend(reps)

    # Dedupe sections by (domain, title_th)
    seen = set()
    deduped_sections = []
    for s in all_sections:
        key = (s["domain"], s["title_th"])
        if key not in seen:
            seen.add(key)
            deduped_sections.append(s)
    if len(deduped_sections) < len(all_sections):
        errors.append(f"Deduped sections: {len(all_sections)} -> {len(deduped_sections)}")

    # Dedupe questions within sections
    for s in deduped_sections:
        seen_q = set()
        deduped_q = []
        for q in s["questions"]:
            norm = re.sub(r"\s+", "", q["text"]).lower()
            key = (norm[:50])
            if key not in seen_q and q["text"]:
                seen_q.add(key)
                deduped_q.append(q)
        if len(deduped_q) < len(s["questions"]):
            errors.append(f"Section '{s['title_th']}': deduped questions {len(s['questions'])} -> {len(deduped_q)}")
        s["questions"] = deduped_q

    # Assign section sequences
    for i, s in enumerate(deduped_sections):
        s["sequence"] = i + 1

    scale_status = SCALE_STATUS.get(layout, "standard_4")

    return {
        "plos": preamble["plos"],
        "sections": deduped_sections,
        "course_codes": preamble["course_codes"],
        "revision_label": preamble["revision_label"],
        "layout": layout,
        "scale_status": scale_status,
        "errors": errors,
        "thai_repairs": all_repairs,
    }


if __name__ == "__main__":
    # Quick test
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else None
    if path:
        result = parse_document(path)
        print(f"Layout: {result['layout']}, scale: {result['scale_status']}")
        print(f"PLOs: {len(result['plos'])}, Sections: {len(result['sections'])}")
        for s in result["sections"]:
            print(f"  {s['domain']} ({s['title_th'][:30]}): {len(s['questions'])} questions")
        if result["errors"]:
            print(f"Errors: {result['errors']}")
