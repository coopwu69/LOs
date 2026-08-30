# -*- coding: utf-8 -*-
"""Convert all source documents (PDF/DOCX/XLSX) to markdown for repo archiving."""
import os, sys, re, glob
import pdfplumber
import docx
import openpyxl

os.chdir(os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding='utf-8')

# Map folder number to short English name (for safe filenames)
FOLDER_MAP = {
    '1.': 'education',
    '2.': 'medicine',
    '3.': 'it',
    '4.': 'political_science',
    '5.': 'science',
    '6.': 'engineering',
    '7.': 'public_health',
    '8.': 'law',
    '9.': 'management',
    '10.': 'accounting',
    '11.': 'architecture',
    '12.': 'international',
    '13.': 'agriculture',
}

OUT_DIR = os.path.join('.', 'docs', 'source')
os.makedirs(OUT_DIR, exist_ok=True)


def get_short_name(folder_name):
    """Get short English name from folder prefix."""
    # Extract leading number
    m = re.match(r'^(\d+)\.', folder_name)
    if not m:
        return 'unknown'
    num = m.group(1) + '.'
    return FOLDER_MAP.get(num, 'unknown')


def convert_pdf(path):
    """Convert PDF to text."""
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            lines.append(text)
    return '\n\n---\n\n'.join(lines)


def convert_docx(path):
    """Convert DOCX to text."""
    d = docx.Document(path)
    lines = []
    for para in d.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    # Also extract tables
    for table in d.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
    return '\n\n'.join(lines)


def convert_xlsx(path):
    """Convert XLSX to text."""
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f'## Sheet: {sheet_name}\n')
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else '' for c in row]
            if any(c.strip() for c in cells):
                lines.append('| ' + ' | '.join(cells) + ' |')
        lines.append('')
    return '\n'.join(lines)


def slugify(text):
    """Make a safe filename from text."""
    # Keep Thai chars, replace spaces and special chars
    text = re.sub(r'[\\/:*?"<>|]', '_', text)
    text = re.sub(r'\s+', '_', text.strip())
    return text[:60]  # Limit length


# Find all source files
extensions = ('*.pdf', '*.docx', '*.xlsx')
all_files = []
for ext in extensions:
    all_files.extend(glob.glob(f'*/**/{ext}', recursive=True))
    all_files.extend(glob.glob(ext))

print(f"Found {len(all_files)} source files")

converted = 0
failed = 0

for filepath in sorted(all_files):
    # Normalize path separators (Windows uses \, we need /)
    filepath = filepath.replace('\\', '/')
    folder = filepath.split('/')[0] if '/' in filepath else ''
    short_name = get_short_name(folder)

    # Build output filename: NN_faculty_filename.md
    folder_num = None
    for prefix in FOLDER_MAP:
        if folder.startswith(prefix):
            # Get the number part
            num_match = re.match(r'^(\d+)', folder)
            if num_match:
                folder_num = int(num_match.group(1))
            break

    basename = os.path.splitext(os.path.basename(filepath))[0]
    safe_base = slugify(basename)

    if folder_num is not None:
        out_name = f'{folder_num:02d}_{short_name}_{safe_base}.md'
    else:
        out_name = f'{short_name}_{safe_base}.md'

    out_path = os.path.join(OUT_DIR, out_name)

    print(f'  Converting: {filepath[:60]}... -> {out_name}', end=' ', flush=True)

    try:
        ext = filepath.lower().rsplit('.', 1)[-1]
        if ext == 'pdf':
            content = convert_pdf(filepath)
        elif ext == 'docx':
            content = convert_docx(filepath)
        elif ext == 'xlsx':
            content = convert_xlsx(filepath)
        else:
            print('SKIP (unknown ext)')
            continue

        # Add header
        header = f'# {basename}\n\n> Source: `{filepath}`\n\n'
        full_content = header + content

        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(full_content)

        print(f'OK ({len(content)} chars)')
        converted += 1
    except Exception as e:
        print(f'FAIL: {str(e)[:80]}')
        failed += 1

print(f'\nDone: {converted} converted, {failed} failed')
print(f'Output: {OUT_DIR}/')
